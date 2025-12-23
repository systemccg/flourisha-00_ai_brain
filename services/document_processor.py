"""
Document Processor Service
Handles processing of various document types with pluggable extraction backends.

Architecture:
- Primary backend: Claude (accurate, handles complex layouts)
- Fallback backend: Docling (free, good for batch processing)
- Validation layer: Catches hallucinations and missing content
"""

import os
import hashlib
import logging
from typing import Dict, Optional, List, Tuple, Any
from pathlib import Path
from enum import Enum

# Legacy imports for backward compatibility
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd

# New extraction backend imports
from .extraction_backends import (
    ExtractionBackend,
    ExtractionResult,
    ClaudeExtractionBackend,
    DoclingExtractionBackend
)
from .extraction_backends.base import ExtractionConfidence, ExtractedEntity

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = {
    'pdf': ['.pdf'],
    'docx': ['.docx', '.doc'],
    'text': ['.txt', '.md', '.markdown'],
    'spreadsheet': ['.xlsx', '.xls', '.csv'],
    'image': ['.png', '.jpg', '.jpeg', '.gif', '.webp']
}

MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB


class ExtractionBackendType(Enum):
    """Available extraction backends"""
    CLAUDE = "claude"       # Primary - accurate, understands context
    DOCLING = "docling"     # Secondary - free, good for batch
    LEGACY = "legacy"       # Original pdfplumber-based extraction


class DocumentProcessor:
    """
    Process and extract text from various document formats.

    Supports pluggable backends:
    - Claude: Best for accuracy and entity extraction
    - Docling: Best for batch processing and simple documents
    - Legacy: Basic extraction using pdfplumber (backward compatible)
    """

    def __init__(
        self,
        default_backend: ExtractionBackendType = ExtractionBackendType.CLAUDE,
        fallback_backend: Optional[ExtractionBackendType] = ExtractionBackendType.DOCLING,
        validate_extractions: bool = True
    ):
        """
        Initialize document processor.

        Args:
            default_backend: Primary extraction backend to use
            fallback_backend: Backend to use if primary fails
            validate_extractions: Whether to validate extraction results
        """
        self.default_backend = default_backend
        self.fallback_backend = fallback_backend
        self.validate_extractions = validate_extractions

        # Lazy-load backends
        self._backends: Dict[ExtractionBackendType, ExtractionBackend] = {}

    def _get_backend(self, backend_type: ExtractionBackendType) -> ExtractionBackend:
        """Get or create backend instance"""
        if backend_type not in self._backends:
            if backend_type == ExtractionBackendType.CLAUDE:
                self._backends[backend_type] = ClaudeExtractionBackend()
            elif backend_type == ExtractionBackendType.DOCLING:
                self._backends[backend_type] = DoclingExtractionBackend()
            else:
                raise ValueError(f"Unknown backend: {backend_type}")

        return self._backends[backend_type]

    @staticmethod
    def detect_file_type(file_path: str) -> Optional[str]:
        """
        Detect file type based on extension

        Args:
            file_path: Path to file

        Returns:
            File type ('pdf', 'docx', 'text', 'spreadsheet', 'image') or None
        """
        ext = Path(file_path).suffix.lower()

        for doc_type, extensions in SUPPORTED_FORMATS.items():
            if ext in extensions:
                return doc_type

        return None

    @staticmethod
    def validate_file(file_path: str) -> Tuple[bool, str]:
        """
        Validate file exists and is within size limit

        Args:
            file_path: Path to file

        Returns:
            (is_valid, message)
        """
        if not os.path.exists(file_path):
            return False, f"File not found: {file_path}"

        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return False, "File is empty"

        if file_size > MAX_FILE_SIZE:
            return False, f"File too large: {file_size} bytes (max: {MAX_FILE_SIZE})"

        return True, "OK"

    async def extract_with_backend(
        self,
        file_path: str,
        backend_type: Optional[ExtractionBackendType] = None,
        extract_entities: bool = True,
        entity_types: Optional[List[str]] = None
    ) -> ExtractionResult:
        """
        Extract document using specified backend with fallback.

        Args:
            file_path: Path to document
            backend_type: Backend to use (defaults to self.default_backend)
            extract_entities: Whether to extract structured entities
            entity_types: Types of entities to extract

        Returns:
            ExtractionResult with content and entities
        """
        backend_type = backend_type or self.default_backend

        # Validate file first
        is_valid, message = self.validate_file(file_path)
        if not is_valid:
            raise ValueError(message)

        # Try primary backend
        try:
            backend = self._get_backend(backend_type)
            result = await backend.extract(
                file_path,
                extract_entities=extract_entities,
                entity_types=entity_types
            )

            # Validate if enabled
            if self.validate_extractions:
                result = await backend.validate_extraction(result, file_path)

            if result.is_valid():
                return result

            # If validation failed, try fallback
            logger.warning(
                f"Primary extraction had issues: {result.validation_errors}. "
                f"Trying fallback backend."
            )

        except Exception as e:
            logger.error(f"Primary backend ({backend_type}) failed: {e}")

            if self.fallback_backend is None:
                raise

        # Try fallback backend
        if self.fallback_backend and self.fallback_backend != backend_type:
            try:
                fallback = self._get_backend(self.fallback_backend)
                result = await fallback.extract(
                    file_path,
                    extract_entities=extract_entities,
                    entity_types=entity_types
                )

                if self.validate_extractions:
                    result = await fallback.validate_extraction(result, file_path)

                result.validation_warnings.append(
                    f"Used fallback backend ({self.fallback_backend.value}) "
                    f"after primary ({backend_type.value}) failed"
                )

                return result

            except Exception as e:
                logger.error(f"Fallback backend also failed: {e}")
                raise

        raise RuntimeError("All extraction backends failed")

    async def process_document(
        self,
        file_path: str,
        use_new_backend: bool = True,
        extract_entities: bool = True,
        entity_types: Optional[List[str]] = None
    ) -> Tuple[str, Dict]:
        """
        Process document and extract text.

        This is the main entry point, backward compatible with legacy usage.

        Args:
            file_path: Path to document file
            use_new_backend: If True, use new pluggable backends; if False, use legacy
            extract_entities: Whether to extract structured entities
            entity_types: Types of entities to extract

        Returns:
            (text, metadata) - Legacy format for backward compatibility

        Raises:
            ValueError: If file format not supported or processing fails
        """
        if use_new_backend:
            result = await self.extract_with_backend(
                file_path,
                extract_entities=extract_entities,
                entity_types=entity_types
            )

            # Convert to legacy format
            metadata = result.metadata.copy()
            metadata['format'] = self.detect_file_type(file_path)
            metadata['backend'] = result.backend_name
            metadata['confidence'] = result.confidence.value
            metadata['entities'] = [e.to_dict() if hasattr(e, 'to_dict') else {
                'name': e.name,
                'type': e.entity_type,
                'value': e.value,
                'confidence': e.confidence.value
            } for e in result.entities]
            metadata['relationships'] = [{
                'source': r.source_entity,
                'target': r.target_entity,
                'type': r.relationship_type
            } for r in result.relationships]
            metadata['validation_warnings'] = result.validation_warnings
            metadata['validation_errors'] = result.validation_errors

            return result.raw_text, metadata

        else:
            # Legacy processing path
            return await self._legacy_process_document(file_path)

    async def _legacy_process_document(self, file_path: str) -> Tuple[str, Dict]:
        """Legacy document processing using pdfplumber/docx/pandas"""
        # Validate file
        is_valid, message = self.validate_file(file_path)
        if not is_valid:
            raise ValueError(message)

        # Detect file type
        file_type = self.detect_file_type(file_path)
        if not file_type:
            raise ValueError(f"Unsupported file format: {Path(file_path).suffix}")

        logger.info(f"Processing {file_type} document (legacy): {file_path}")

        # Process based on type
        if file_type == 'pdf':
            return await self._extract_pdf_text_legacy(file_path)
        elif file_type == 'docx':
            return await self._extract_docx_text(file_path)
        elif file_type == 'text':
            return await self._extract_text_file(file_path)
        elif file_type == 'spreadsheet':
            return await self._extract_spreadsheet_data(file_path)
        else:
            raise ValueError(f"Unknown file type: {file_type}")

    # =====================================================
    # Legacy extraction methods (backward compatibility)
    # =====================================================

    @staticmethod
    async def _extract_pdf_text_legacy(file_path: str) -> Tuple[str, Dict]:
        """Legacy PDF extraction using pdfplumber"""
        text = ""
        metadata = {
            'format': 'pdf',
            'pages': 0,
            'images': 0,
            'tables': 0,
            'backend': 'legacy'
        }

        try:
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)

                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"

                    tables = page.extract_tables()
                    if tables:
                        metadata['tables'] += len(tables)
                        for table in tables:
                            for row in table:
                                text += " | ".join(str(cell) for cell in row if cell) + "\n"

            logger.info(f"Extracted {metadata['pages']} pages from PDF (legacy)")
            return text, metadata

        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise

    @staticmethod
    async def _extract_docx_text(file_path: str) -> Tuple[str, Dict]:
        """Extract text from DOCX file"""
        text = ""
        metadata = {
            'format': 'docx',
            'paragraphs': 0,
            'tables': 0
        }

        try:
            doc = DocxDocument(file_path)

            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
                    metadata['paragraphs'] += 1

            for table in doc.tables:
                metadata['tables'] += 1
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        text += row_text + "\n"

            logger.info(f"Extracted {metadata['paragraphs']} paragraphs from DOCX")
            return text, metadata

        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise

    @staticmethod
    async def _extract_text_file(file_path: str) -> Tuple[str, Dict]:
        """Extract text from plain text file"""
        metadata = {
            'format': 'text',
            'lines': 0,
            'encoding': 'utf-8'
        }

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                metadata['lines'] = len(text.split('\n'))

            logger.info(f"Extracted {metadata['lines']} lines from text file")
            return text, metadata

        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
                metadata['lines'] = len(text.split('\n'))
                metadata['encoding'] = 'latin-1'

            logger.info(f"Extracted {metadata['lines']} lines from text file (latin-1)")
            return text, metadata

        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            raise

    @staticmethod
    async def _extract_spreadsheet_data(file_path: str) -> Tuple[str, Dict]:
        """Extract data from spreadsheet (Excel, CSV)"""
        metadata = {
            'format': 'spreadsheet',
            'sheets': 0,
            'rows': 0,
            'columns': 0
        }

        text = ""

        try:
            ext = Path(file_path).suffix.lower()

            if ext == '.csv':
                df = pd.read_csv(file_path)
                sheets = {'Sheet1': df}
                metadata['sheets'] = 1
            else:
                xls = pd.ExcelFile(file_path)
                sheets = {}
                for sheet_name in xls.sheet_names:
                    sheets[sheet_name] = pd.read_excel(file_path, sheet_name=sheet_name)
                metadata['sheets'] = len(sheets)

            for sheet_name, df in sheets.items():
                text += f"\n=== Sheet: {sheet_name} ===\n"
                text += df.to_string(index=False)
                text += "\n"

                metadata['rows'] += len(df)
                metadata['columns'] = max(metadata['columns'], len(df.columns))

            logger.info(f"Extracted {metadata['sheets']} sheets, {metadata['rows']} rows")
            return text, metadata

        except Exception as e:
            logger.error(f"Spreadsheet extraction error: {str(e)}")
            raise

    # =====================================================
    # Utility methods
    # =====================================================

    @staticmethod
    def calculate_content_hash(content: str) -> str:
        """Calculate SHA-256 hash of content"""
        return hashlib.sha256(content.encode()).hexdigest()

    @staticmethod
    def extract_summary_stats(text: str) -> Dict:
        """Extract basic statistics about text"""
        words = text.split()
        lines = text.split('\n')

        return {
            'character_count': len(text),
            'word_count': len(words),
            'line_count': len(lines),
            'average_word_length': len(text) / len(words) if words else 0
        }


# Backward-compatible aliases
extract_pdf_text = DocumentProcessor._extract_pdf_text_legacy
extract_docx_text = DocumentProcessor._extract_docx_text
extract_text_file = DocumentProcessor._extract_text_file
extract_spreadsheet_data = DocumentProcessor._extract_spreadsheet_data

# Global instance
_processor = None


async def get_document_processor(
    default_backend: ExtractionBackendType = ExtractionBackendType.CLAUDE,
    fallback_backend: Optional[ExtractionBackendType] = ExtractionBackendType.DOCLING
) -> DocumentProcessor:
    """Get or create document processor instance"""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor(
            default_backend=default_backend,
            fallback_backend=fallback_backend
        )
    return _processor


async def process_file(
    file_path: str,
    use_new_backend: bool = True,
    extract_entities: bool = True
) -> Tuple[str, Dict]:
    """
    Convenience function to process a file

    Args:
        file_path: Path to file
        use_new_backend: Use new pluggable backends (Claude/Docling)
        extract_entities: Extract structured entities

    Returns:
        (text, metadata)
    """
    processor = await get_document_processor()
    return await processor.process_document(
        file_path,
        use_new_backend=use_new_backend,
        extract_entities=extract_entities
    )
